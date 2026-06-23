"""
Сервис генерации документов:
- PDF реестр (Таблица 3) + паспорта
- Word реестр
- PDF паспорт одной РК
"""
import io
from pathlib import Path
from typing import Optional
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph,
    Spacer, Image, PageBreak, KeepTogether
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import docx.oxml as oxml

from app.models.rk import RK
from app.config import settings


# ─── Константы ────────────────────────────────────────────────────────────────

REGISTRY_COLUMNS = [
    ("№ п/п", 0.8*cm),
    ("№ РК", 1.2*cm),
    ("Местоположение", 4.5*cm),
    ("Тип", 2.0*cm),
    ("Вид РК", 2.5*cm),
    ("Размер (м)", 2.0*cm),
    ("Площадь (м²)", 1.5*cm),
    ("Координаты WGS84", 3.5*cm),
    ("Соответствие", 2.5*cm),
    ("Правоустан. документ", 2.0*cm),
    ("Примечание", 2.5*cm),
]

BRAND_BLUE = colors.HexColor("#1B4F8A")
HEADER_GRAY = colors.HexColor("#E8F0F8")

TYPE_ORDER = [
    "Сити-формат", "Билборд динамика", "Цифровой билборд",
    "Панель-кронштейн", "Остановочный пункт", "Билборд",
    "Настенный щит", "Афиша", "Прочее"
]


# ─── PDF Реестр ───────────────────────────────────────────────────────────────

def generate_registry_pdf(rks: list[RK]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=2*cm, bottomMargin=1.5*cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Normal"],
        fontSize=13, fontName="Helvetica-Bold",
        alignment=TA_CENTER, spaceAfter=4
    )
    sub_style = ParagraphStyle(
        "Sub", parent=styles["Normal"],
        fontSize=11, fontName="Helvetica",
        alignment=TA_CENTER, spaceAfter=20
    )
    cell_style = ParagraphStyle(
        "Cell", parent=styles["Normal"],
        fontSize=7, fontName="Helvetica",
        leading=9
    )
    cell_bold = ParagraphStyle(
        "CellBold", parent=cell_style,
        fontName="Helvetica-Bold"
    )

    story = []
    story.append(Paragraph("Таблица 3", sub_style))
    story.append(Paragraph(
        "РЕЕСТР рекламных конструкций на территории<br/>"
        "Шпаковского муниципального округа Ставропольского края",
        title_style
    ))
    story.append(Spacer(1, 0.5*cm))

    # Группировка по типу
    grouped: dict[str, list[RK]] = {}
    for rk in rks:
        grouped.setdefault(rk.type_rk, []).append(rk)

    col_widths = [w for _, w in REGISTRY_COLUMNS]

    for type_name in TYPE_ORDER:
        group = grouped.get(type_name, [])
        if not group:
            continue

        # Заголовок группы
        group_header = [[
            Paragraph(type_name, cell_bold),
            *[""] * (len(REGISTRY_COLUMNS) - 1)
        ]]
        group_table = Table(group_header, colWidths=col_widths)
        group_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HEADER_GRAY),
            ("SPAN",       (0, 0), (-1, 0)),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, 0), 8),
            ("ALIGN",      (0, 0), (-1, 0), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, 0), 4),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.gray),
        ]))
        story.append(group_table)

        # Строки данных
        table_data = [[
            Paragraph(h, cell_bold) for h, _ in REGISTRY_COLUMNS
        ]]

        for i, rk in enumerate(group):
            coords = f"{rk.lat:.8f},\n{rk.lon:.8f}" if rk.lat else "—"
            table_data.append([
                Paragraph(str(rk.num),        cell_style),
                Paragraph(rk.rk_id,           cell_bold),
                Paragraph(rk.address or "—",  cell_style),
                Paragraph(rk.type_adv or "—", cell_style),
                Paragraph(rk.type_rk,         cell_style),
                Paragraph(rk.size or "—",     cell_style),
                Paragraph(rk.area or "—",     cell_style),
                Paragraph(coords,             cell_style),
                Paragraph(rk.compliance or "—", cell_style),
                Paragraph(rk.legal_doc or "—", cell_style),
                Paragraph(rk.note or "—",     cell_style),
            ])

        data_table = Table(table_data, colWidths=col_widths, repeatRows=1)
        data_table.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), BRAND_BLUE),
            ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, 0), 7),
            ("ALIGN",         (0, 0), (-1, 0), "CENTER"),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING",   (0, 0), (-1, -1), 3),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 3),
            ("GRID",          (0, 0), (-1, -1), 0.4, colors.gray),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ]))
        story.append(data_table)
        story.append(Spacer(1, 0.4*cm))

    # Нижний колонтитул
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        f"Итого: {len(rks)} рекламных конструкций  |  "
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}",
        ParagraphStyle("footer", parent=styles["Normal"],
                       fontSize=8, alignment=TA_CENTER,
                       textColor=colors.gray)
    ))

    doc.build(story)
    return buf.getvalue()


# ─── PDF Паспорт одной РК ─────────────────────────────────────────────────────

def generate_passport_pdf(rk: RK) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", fontSize=13, fontName="Helvetica-Bold",
                         alignment=TA_LEFT, spaceAfter=14)
    label = ParagraphStyle("label", fontSize=9, fontName="Helvetica-Bold",
                            textColor=colors.HexColor("#475569"))
    value = ParagraphStyle("value", fontSize=10, fontName="Helvetica")

    story = []
    story.append(Paragraph(f"Рекламная конструкция № {rk.rk_id}", h1))

    # Таблица с данными
    info_data = [
        ["Тип рекламной конструкции", rk.type_rk or "—"],
        ["Адрес", rk.address or "—"],
        ["Координаты МСК", f"{rk.msk_x}  {rk.msk_y}" if rk.msk_x else "—"],
        ["Координаты WGS84", f"{rk.lat:.8f}    {rk.lon:.8f}" if rk.lat else "—"],
        ["Размер", rk.size or "—"],
        ["Общая площадь, м²", rk.area or "—"],
        ["Соответствие", rk.compliance or "—"],
        ["Правоустанавл. документ", rk.legal_doc or "—"],
        ["Примечание", rk.note or "—"],
    ]

    info_table = Table(
        [[Paragraph(r[0], label), Paragraph(r[1], value)] for r in info_data],
        colWidths=[5.5*cm, 12*cm]
    )
    info_table.setStyle(TableStyle([
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#E2E8F0")),
        ("BACKGROUND",    (0, 0), (0, -1), colors.HexColor("#F8FAFC")),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.6*cm))

    # Схема расположения
    story.append(Paragraph("Схема расположения", label))
    story.append(Spacer(1, 0.2*cm))
    if rk.scheme_path and Path(rk.scheme_path).exists():
        try:
            img = Image(rk.scheme_path, width=16*cm, height=9*cm, kind="proportional")
            story.append(img)
        except Exception:
            story.append(Paragraph("[Схема недоступна]", value))
    else:
        placeholder = Table([["Схема расположения отсутствует"]], colWidths=[16*cm])
        placeholder.setStyle(TableStyle([
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
            ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor("#F1F5F9")),
            ("TOPPADDING",    (0, 0), (-1, -1), 40),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 40),
            ("TEXTCOLOR",     (0, 0), (-1, -1), colors.HexColor("#94A3B8")),
        ]))
        story.append(placeholder)

    story.append(Spacer(1, 0.6*cm))

    # Фото объекта
    story.append(Paragraph("Фото", label))
    story.append(Spacer(1, 0.2*cm))
    if rk.photo_path and Path(rk.photo_path).exists():
        try:
            img = Image(rk.photo_path, width=16*cm, height=9*cm, kind="proportional")
            story.append(img)
        except Exception:
            story.append(Paragraph("[Фото недоступно]", value))
    else:
        placeholder = Table([["Фото отсутствует"]], colWidths=[16*cm])
        placeholder.setStyle(TableStyle([
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
            ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor("#F1F5F9")),
            ("TOPPADDING",    (0, 0), (-1, -1), 40),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 40),
            ("TEXTCOLOR",     (0, 0), (-1, -1), colors.HexColor("#94A3B8")),
        ]))
        story.append(placeholder)

    doc.build(story)
    return buf.getvalue()


# ─── Word Реестр ──────────────────────────────────────────────────────────────

def generate_registry_docx(rks: list[RK]) -> bytes:
    doc = Document()

    # Страница — альбомная
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(1.5)

    # Заголовок
    t = doc.add_paragraph("Таблица 3")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(11)

    h = doc.add_paragraph(
        "РЕЕСТР рекламных конструкций на территории\n"
        "Шпаковского муниципального округа Ставропольского края"
    )
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(13)
    doc.add_paragraph()

    # Заголовки столбцов
    headers = [h for h, _ in REGISTRY_COLUMNS]

    # Группировка
    grouped: dict[str, list[RK]] = {}
    for rk in rks:
        grouped.setdefault(rk.type_rk, []).append(rk)

    for type_name in TYPE_ORDER:
        group = grouped.get(type_name, [])
        if not group:
            continue

        # Заголовок группы
        gp = doc.add_paragraph(type_name)
        gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gp.runs[0].font.bold = True
        gp.runs[0].font.size = Pt(10)
        gp.runs[0].font.color.rgb = RGBColor(0x1B, 0x4F, 0x8A)

        # Таблица группы
        table = doc.add_table(rows=1 + len(group), cols=len(headers))
        table.style = "Table Grid"

        # Заголовочная строка
        hrow = table.rows[0]
        for i, header in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = oxml.OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1B4F8A")
            tcPr.append(shd)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Строки данных
        for row_idx, rk in enumerate(group):
            coords = f"{rk.lat:.8f}, {rk.lon:.8f}" if rk.lat else "—"
            values = [
                str(rk.num),
                rk.rk_id,
                rk.address or "—",
                rk.type_adv or "—",
                rk.type_rk,
                rk.size or "—",
                rk.area or "—",
                coords,
                rk.compliance or "—",
                rk.legal_doc or "—",
                rk.note or "—",
            ]
            row = table.rows[row_idx + 1]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(7)

                # Чередование строк
                if row_idx % 2 == 1:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = oxml.OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F8FAFC")
                    tcPr.append(shd)

        doc.add_paragraph()

    # Итог
    footer = doc.add_paragraph(
        f"Итого: {len(rks)} рекламных конструкций  |  "
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
